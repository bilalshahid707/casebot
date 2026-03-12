from pydantic import BaseModel
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from agents.ocr_agent.agent import OCRAgent
import io


class ParsedDocument(BaseModel):
    page_content: str
    source: str
    page_number: int | None = None


def parse_file_to_documents(file: dict) -> list[ParsedDocument]:

    documents: list[ParsedDocument] = []
    content_type = file["content_type"]
    source = file["filename"]

    if content_type.startswith("image/"):
        text = OCRAgent().extract_text(file)
        documents.append(ParsedDocument(page_content=text, source=source))

    elif content_type == "application/pdf":
        pdf = PdfReader(io.BytesIO(file["content"]))
        for i, page in enumerate(pdf.pages):
            documents.append(
                ParsedDocument(
                    page_content=page.extract_text() or "",
                    source=source,
                    page_number=i + 1,
                )
            )

    elif content_type in (
        "application/msword",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ):
        docx = DocxDocument(io.BytesIO(file["content"]))
        for para in docx.paragraphs:
            if para.text.strip():
                documents.append(ParsedDocument(page_content=para.text, source=source))

    elif content_type in ("text/plain", "text/markdown"):
        text = file["content"].decode("utf-8")
        documents.append(ParsedDocument(page_content=text, source=source))

    else:
        raise ValueError(f"Unsupported file type: {content_type}")

    return documents
