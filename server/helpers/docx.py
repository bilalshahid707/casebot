from docx import Document
import io


def convert_text_to_docx(text_model):
    text_dict = text_model.model_dump()
    doc = Document()

    for key, value in text_dict.items():
        doc.add_heading(str(key), level=1)
        doc.add_paragraph(str(value))

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream
